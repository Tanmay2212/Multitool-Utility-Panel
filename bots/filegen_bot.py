import os
import tkinter as tk
from tkinter import filedialog, messagebox
from utils.logger import log_event

class FileGenUI:
    def __init__(self, root, parent_frame):
        self.frame = tk.Frame(parent_frame, bg="#ecf0f1", padx=20, pady=20)
        self.frame.pack(fill="both", expand=True)

        tk.Label(self.frame, text="🧪 Dummy File Generator", font=("Arial", 14, "bold"), bg="#ecf0f1").pack(anchor="w", pady=(0, 15))

        # File type selector
        tk.Label(self.frame, text="Select File Type:", bg="#ecf0f1").pack(anchor="w")
        self.file_type = tk.StringVar(value="txt")
        options = ["txt", "pdf", "docx", "xlsx"]
        tk.OptionMenu(self.frame, self.file_type, *options).pack(anchor="w", pady=(0, 10))

        # File name
        tk.Label(self.frame, text="Enter File Name:", bg="#ecf0f1").pack(anchor="w")
        self.file_name_entry = tk.Entry(self.frame, width=50)
        self.file_name_entry.pack(anchor="w", pady=(0, 10))

        # Output folder
        tk.Label(self.frame, text="Save Location:", bg="#ecf0f1").pack(anchor="w")
        self.folder_var = tk.StringVar(value=os.path.expanduser("~/Downloads"))
        tk.Entry(self.frame, textvariable=self.folder_var, width=50).pack(anchor="w")
        tk.Button(self.frame, text="📁 Browse", command=self.browse_folder).pack(anchor="w", pady=(5, 10))

        # Create file button
        tk.Button(self.frame, text="⚙️ Generate File", command=self.generate_file, bg="#2ecc71", fg="white").pack(anchor="w", pady=10)

    def browse_folder(self):
        folder = filedialog.askdirectory()
        if folder:
            self.folder_var.set(folder)

    def generate_file(self):
        name = self.file_name_entry.get().strip()
        folder = self.folder_var.get().strip()
        ext = self.file_type.get()

        if not name:
            messagebox.showwarning("Input Needed", "Please enter a file name.")
            return

        file_path = os.path.join(folder, f"{name}.{ext}")

        try:
            if ext == "txt":
                with open(file_path, "w") as f:
                    f.write("This is a dummy text file.\nJust for testing.")
            elif ext == "pdf":
                from fpdf import FPDF
                pdf = FPDF()
                pdf.add_page()
                pdf.set_font("Arial", size=12)
                pdf.cell(200, 10, txt="Dummy PDF File", ln=True)
                pdf.output(file_path)
            elif ext == "docx":
                from docx import Document
                doc = Document()
                doc.add_heading("Dummy DOCX File", 0)
                doc.add_paragraph("This is a test document.")
                doc.save(file_path)
            elif ext == "xlsx":
                import openpyxl
                wb = openpyxl.Workbook()
                ws = wb.active
                ws["A1"] = "Dummy Excel File"
                wb.save(file_path)
            else:
                raise ValueError("Unsupported file type.")

            messagebox.showinfo("Success", f"File created:\n{file_path}")
            log_event("File Generator", f"Generated {ext} file: {file_path}")

        except Exception as e:
            messagebox.showerror("Error", f"Something went wrong: {e}")
